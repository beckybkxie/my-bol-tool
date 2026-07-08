"""
FBA BOL 自动化生成系统
------------------------------------------
数据库：Supabase (Postgres)，warehouses 表以 name 为主键，
        每个仓库代码只保留一条记录，重复添加会覆盖旧地址。
部署前请在 Supabase 中执行以下 SQL 建表：

    create table warehouses (
      name text primary key,
      address text not null,
      updated_at timestamp with time zone default now()
    );

并在 .streamlit/secrets.toml（本地）或 Streamlit Cloud 的 Secrets 设置中配置：

    SUPABASE_URL = "https://你的项目.supabase.co"
    SUPABASE_KEY = "你的anon_key"

依赖安装：
    pip install streamlit pandas python-docx supabase
"""

import streamlit as st
import pandas as pd
from docx import Document
from datetime import datetime
import io
import os
from supabase import create_client

st.set_page_config(page_title="BOL 自动化系统", layout="wide")

# ============================================================
# Supabase 连接与数据操作
# ============================================================

@st.cache_resource
def get_supabase():
    """建立到 Supabase 的连接，全局复用（不会每次刷新都重连）。"""
    url = st.secrets["SUPABASE_URL"]
    key = st.secrets["SUPABASE_KEY"]
    return create_client(url, key)


def load_warehouses():
    """从数据库读取全部仓库记录，按名称排序。返回列表，每条为 dict。"""
    supabase = get_supabase()
    res = supabase.table("warehouses").select("*").order("name").execute()
    return res.data


def upsert_warehouse(name: str, address: str):
    """
    新增或更新仓库地址。
    name 是主键：已存在则覆盖 address，不存在则新增一条。
    """
    supabase = get_supabase()
    supabase.table("warehouses").upsert(
        {"name": name, "address": address}
    ).execute()


def delete_warehouse(name: str):
    """按名称删除一条仓库记录。"""
    supabase = get_supabase()
    supabase.table("warehouses").delete().eq("name", name).execute()


# ============================================================
# BOL 文档生成逻辑
# ============================================================

def generate_bol(data_rows, ship_to_addr, isa_no, full_appt, selected_wh):
    """
    读取本地模板 bol_template.docx，替换占位符并填充 PRO 表格数据，
    返回内存中的 docx 文件流和生成的 BOL 编号。
    """
    template_path = "bol_template.docx"
    if not os.path.exists(template_path):
        return None, "找不到模板文件 bol_template.docx，请确认它与本脚本放在同一目录下"

    doc = Document(template_path)
    today_str = datetime.now().strftime("%Y-%m-%d")
    bol_number = "BOL" + datetime.now().strftime("%Y%m%d%H%M%S") + "-" + selected_wh

    placeholders = {
        "{{DATE}}": today_str,
        "{{BOL_NO}}": bol_number,
        "{{SHIP_TO}}": f"{selected_wh} - {ship_to_addr}",
        "{{ISA}}": isa_no,
        "{{APPOINTMENT}}": full_appt,
    }

    # 替换段落 + 表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

    # 找到包含 "PRO" 表头的表格，清空旧数据行，写入新数据
    for table in doc.tables:
        if len(table.rows) > 0 and "PRO" in table.rows[0].cells[0].text.upper():
            while len(table.rows) > 1:
                row_to_remove = table.rows[-1]._element
                row_to_remove.getparent().remove(row_to_remove)
            for row_data in data_rows:
                cells = table.add_row().cells
                for i in range(min(len(cells), len(row_data))):
                    cells[i].text = str(row_data[i])
            break

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio, bol_number


# ============================================================
# 侧边栏：仓库地址数据库管理（Supabase）
# ============================================================

warehouses = load_warehouses()
warehouse_dict = {w["name"]: w["address"] for w in warehouses}

with st.sidebar:
    st.header("🏢 仓库数据库管理")
    st.caption("数据保存在云端数据库，所有登录本 App 的同事共享同一份地址库。")

    with st.expander("➕ 添加 / 修改地址", expanded=False):
        new_name = st.text_input("仓库代码 (如 LAX9)")
        new_addr = st.text_area("详细地址")

        is_existing = new_name in warehouse_dict
        if new_name and is_existing:
            st.info(f"⚠️ 仓库代码「{new_name}」已存在，保存将覆盖原地址：\n\n{warehouse_dict[new_name]}")

        if st.button("💾 保存到数据库"):
            if new_name and new_addr:
                upsert_warehouse(new_name, new_addr)
                if is_existing:
                    st.success(f"已覆盖更新「{new_name}」的地址")
                else:
                    st.success(f"已新增仓库「{new_name}」")
                st.rerun()
            else:
                st.error("请填写完整的仓库代码和地址")

    with st.expander("🗑️ 查看 / 删除已有地址", expanded=False):
        if warehouse_dict:
            for w in warehouses:
                st.markdown(f"**{w['name']}**  \n{w['address']}")
            st.divider()
            to_delete = st.selectbox("选择要删除的仓库代码", list(warehouse_dict.keys()))
            if st.button("确认删除"):
                delete_warehouse(to_delete)
                st.success(f"已删除「{to_delete}」")
                st.rerun()
        else:
            st.info("数据库中暂无仓库地址，请先添加")


# ============================================================
# 主界面：上传 PO -> 确认信息 -> 生成 BOL
# ============================================================

st.title("📦 FBA BOL 自动化生成系统")

st.header("1️⃣ 上传 PO 文件")
uploaded_excel = st.file_uploader("选择 Excel 文件", type=["xlsx", "xls"])

if uploaded_excel:
    df_raw = pd.read_excel(uploaded_excel, header=None)
    extracted_df = df_raw.iloc[1:, [1, 3, 5, 7]].copy()
    extracted_df.dropna(how="all", inplace=True)

    st.subheader("提取到的数据预览")
    st.dataframe(extracted_df.head())

    st.header("2️⃣ 确认派送信息")

    col1, col2 = st.columns(2)
    with col1:
        if warehouse_dict:
            selected_wh = st.selectbox("选择目标仓库", list(warehouse_dict.keys()))
            ship_to_addr = warehouse_dict[selected_wh]
            st.caption(f"📍 当前地址：{selected_wh} - {ship_to_addr}")
        else:
            st.warning("请先在左侧边栏「仓库数据库管理」中添加至少一个仓库地址")
            selected_wh, ship_to_addr = "", ""

        isa_no = st.text_input("ISA#")

    with col2:
        appt_date = st.date_input("预约日期")
        appt_time = st.time_input("预约时间")
        full_appt = f"{appt_date} {appt_time}"

    st.header("3️⃣ 生成文件")
    if st.button("🚀 生成并下载 BOL"):
        if not selected_wh:
            st.error("请先选择目标仓库")
        else:
            result, b_no = generate_bol(
                extracted_df.values.tolist(),
                ship_to_addr,
                isa_no,
                full_appt,
                selected_wh,
            )
            if result:
                st.success(f"生成成功：{b_no}")
                st.download_button(
                    "📥 点击下载 BOL 文件",
                    result,
                    f"BOL_{b_no}.docx",
                )
            else:
                st.error(b_no)
else:
    st.info("请上传 PO Excel 文件以开始")
