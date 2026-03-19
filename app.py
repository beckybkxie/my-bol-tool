import streamlit as st
import pandas as pd
from docx import Document
from datetime import datetime
import io
import os
import json

st.set_page_config(page_title="BOL 自动化系统", layout="wide")

# --- 数据库操作：仓库地址管理 ---
DB_FILE = "warehouses.json"

def load_warehouses():
    if os.path.exists(DB_FILE):
        with open(DB_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"Amazon - ONT8": "24300 Nandina Ave, Moreno Valley, CA 92551"}

def save_warehouses(data):
    with open(DB_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)

# 初始化数据
if 'warehouses' not in st.session_state:
    st.session_state.warehouses = load_warehouses()

# --- 侧边栏：地址数据库管理 ---
with st.sidebar:
    st.header("🏢 仓库数据库管理")
    with st.expander("添加/修改地址"):
        new_name = st.text_input("仓库名称 (如 LAX9)")
        new_addr = st.text_area("详细地址")
        if st.button("保存到数据库"):
            if new_name and new_addr:
                st.session_state.warehouses[new_name] = new_addr
                save_warehouses(st.session_state.warehouses)
                st.success(f"已更新 {new_name}")
            else:
                st.error("请填写完整信息")
    
    if st.checkbox("查看/删除已有地址"):
        to_delete = st.selectbox("选择要删除的地址", list(st.session_state.warehouses.keys()))
        if st.button("确认删除"):
            del st.session_state.warehouses[to_delete]
            save_warehouses(st.session_state.warehouses)
            st.rerun()

# --- 核心生成逻辑 (保持不变，仅修改 ship_to 来源) ---
def generate_bol(data_rows, ship_to_addr, isa_no, appt_info):
    template_path = "bol_template.docx"
    if not os.path.exists(template_path):
        return None, "找不到模板文件"
    
    doc = Document(template_path)
    today_str = datetime.now().strftime("%Y-%m-%d")
    bol_number = "BOL" + datetime.now().strftime("%Y%m%d%H%M%S")
    
    placeholders = {
        "{{DATE}}": today_str,
        "{{BOL_NO}}": bol_number,
        "{{SHIP_TO}}": ship_to_addr,
        "{{ISA}}": isa_no,
        "{{APPOINTMENT}}": appt_info
    }

    # 替换逻辑 (段落+表格)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))
    
    # 填充 B,D,F,H 列数据到含有 PRO 的表
    for table in doc.tables:
        if len(table.rows) > 0 and "PRO" in table.rows[0].cells[0].text.upper():
            while len(table.rows) > 1:
                row_to_remove = table.rows[-1]._element
                row_to_remove.getparent().remove(row_to_remove)
            for row_data in data_rows:
                cells = table.add_row().cells
                for i in range(min(len(cells), len(row_data))):
                    cells[i].text = str(row_data[i])
            break

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio, bol_number

# --- 主界面 ---
st.header("🤨1. 上传 PO")
uploaded_excel = st.file_uploader("选择文件", type=["xlsx", "xls"])

if uploaded_excel:
    df_raw = pd.read_excel(uploaded_excel, header=None)
    extracted_df = df_raw.iloc[1:, [1, 3, 5, 7]].copy()
    extracted_df.dropna(how='all', inplace=True)
    st.dataframe(extracted_df.head())

    st.header("2. 确认派送信息")
    
    col1, col2 = st.columns(2)
    with col1:
        # 从“数据库”中选择
        selected_wh = st.selectbox("选择目标仓库", list(st.session_state.warehouses.keys()))
        ship_to_addr = st.session_state.warehouses[selected_wh]
        st.caption(f"当前地址: {selected_wh}: {ship_to_addr}")
        
        isa_no = st.text_input("ISA#")
    with col2:
        appt_date = st.date_input("预约日期")
        appt_time = st.time_input("预约时间")
        full_appt = f"{appt_date} {appt_time}"

    if st.button("🚀 生成并下载"):
        result, b_no = generate_bol(extracted_df.values.tolist(), ship_to_addr, isa_no, full_appt)
        if result:
            st.download_button("📥 点击下载", result, f"BOL_{b_no}.docx")